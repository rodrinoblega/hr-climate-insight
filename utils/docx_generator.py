"""
DOCX generator - converts Markdown to Word document with chart support.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Optional


def markdown_to_docx(
    markdown_text: str,
    output_path: str | Path,
    charts: Optional[dict] = None,
) -> Path:
    """
    Convert Markdown text to a formatted Word document.

    Args:
        markdown_text: The markdown content to convert
        output_path: Path for the output DOCX file
        charts: Dictionary mapping question numbers/text to chart info
                Format: {question_key: {'path': Path, 'info': dict}}

    Returns:
        Path to the generated DOCX file
    """
    output_path = Path(output_path)
    doc = Document()

    # Set up styles
    _setup_styles(doc)

    # Process markdown line by line
    lines = markdown_text.split('\n')
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Handle chart markers: [GRAFICO: ...] or [CHART: ...]
        chart_match = re.match(r'^\s*\[(GRAFICO|CHART|GRÁFICO):\s*(.+?)\]\s*$', line, re.IGNORECASE)
        if chart_match and charts:
            chart_key = chart_match.group(2).strip()
            _insert_chart(doc, chart_key, charts)
            i += 1
            continue

        # Handle tables
        if '|' in line and not line.strip().startswith('```'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            _add_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)

        # Handle horizontal rules
        elif line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 50)

        # Handle bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            text = _process_inline_formatting(text)
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = _process_inline_formatting(text)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)

        # Handle bold lines (like **Para RRHH:**)
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip()[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True

        # Handle regular paragraphs
        elif line.strip():
            text = _process_inline_formatting(line)
            p = doc.add_paragraph()
            _add_formatted_text(p, text)

        i += 1

    # Handle any remaining table
    if in_table and table_rows:
        _add_table(doc, table_rows)

    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return output_path


def _insert_chart(doc: Document, chart_key: str, charts: dict):
    """
    Insert a chart image into the document.

    Args:
        doc: Document object
        chart_key: Keyword to look up the chart (must match key in charts dict)
        charts: Dictionary from generate_all_charts() with keywords as keys
    """
    chart_key_clean = chart_key.lower().strip()

    # 1. Direct lookup by keyword
    chart_info = charts.get(chart_key_clean)

    # 2. Fallback: case-insensitive exact match
    if not chart_info:
        for key, info in charts.items():
            if key.lower() == chart_key_clean:
                chart_info = info
                break

    # 3. Fallback: partial match (keyword contains search term or vice versa)
    if not chart_info:
        for key, info in charts.items():
            key_lower = key.lower()
            if chart_key_clean in key_lower or key_lower in chart_key_clean:
                chart_info = info
                break

    # 4. Fallback: search in question text
    if not chart_info:
        for key, info in charts.items():
            question = info.get('question', '').lower()
            if chart_key_clean in question:
                chart_info = info
                break

    # 5. Fallback: related terms mapping
    if not chart_info:
        related_terms = {
            'liderazgo': ['lider', 'superior', 'jefe', 'feedback', 'escucha'],
            'comunicacion': ['comunicación', 'informa', 'transparencia'],
            'equipo': ['equipo_trabajo', 'colaboracion', 'compañeros'],
            'compensacion': ['remuneracion', 'salario', 'sueldo', 'beneficios'],
            'desarrollo': ['capacitacion', 'formacion', 'crecimiento'],
            'compromiso': ['orgullo', 'pertenencia', 'recomendar'],
            'clima': ['clima_laboral', 'ambiente'],
        }
        for term, related in related_terms.items():
            if chart_key_clean == term or chart_key_clean in related:
                # Search for any of the related terms
                for key, info in charts.items():
                    key_lower = key.lower()
                    if key_lower == term or key_lower in related or term in key_lower:
                        chart_info = info
                        break
                if chart_info:
                    break

    if chart_info and chart_info.get('path') and Path(chart_info['path']).exists():
        # Add the image centered
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(chart_info['path']), width=Inches(5.5))

        # Add a small space after the image
        doc.add_paragraph()
    else:
        # Chart not found - add a placeholder note
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Gráfico no disponible: {chart_key}]")
        run.italic = True


def _setup_styles(doc: Document):
    """Set up document styles."""
    # Modify Normal style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


def _process_inline_formatting(text: str) -> str:
    """Process inline markdown formatting markers."""
    return text


def _add_formatted_text(paragraph, text: str):
    """Add text with inline formatting to a paragraph."""
    # Pattern to find **bold** and *italic* text
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _add_table(doc: Document, table_rows: list[str]):
    """Add a markdown table to the document."""
    if len(table_rows) < 2:
        return

    # Parse table rows
    parsed_rows = []
    for row in table_rows:
        # Skip separator row (|---|---|)
        if re.match(r'^[\s|:-]+$', row):
            continue
        cells = [cell.strip() for cell in row.split('|')]
        # Remove empty first/last cells from | delimited rows
        cells = [c for c in cells if c]
        if cells:
            parsed_rows.append(cells)

    if not parsed_rows:
        return

    # Create table
    num_cols = len(parsed_rows[0])
    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    table.style = 'Table Grid'

    # Fill table
    for i, row_data in enumerate(parsed_rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                # Remove markdown formatting for table cells
                clean_text = re.sub(r'\*+', '', cell_text)
                row.cells[j].text = clean_text

    # Bold header row
    if table.rows:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    doc.add_paragraph()  # Add spacing after table


def save_markdown(markdown_text: str, output_path: str | Path) -> Path:
    """
    Save the raw markdown to a file (useful for debugging).

    Args:
        markdown_text: The markdown content
        output_path: Path for the output file

    Returns:
        Path to the saved file
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(markdown_text, encoding='utf-8')
    return output_path
